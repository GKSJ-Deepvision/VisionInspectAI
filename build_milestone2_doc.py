import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_callout(doc, text, title="NOTE", fill_hex="F0F4F8", border_hex="0078D4"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, fill_hex)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    
    run_title = p.add_run(f"[{title}] ")
    run_title.bold = True
    run_title.font.name = 'Segoe UI'
    run_title.font.size = Pt(10.5)
    run_title.font.color.rgb = RGBColor(0, 80, 160)
    
    run_text = p.add_run(text)
    run_text.font.name = 'Segoe UI'
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = RGBColor(40, 40, 40)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def build_docx():
    doc = docx.Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Set base normal style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(30, 30, 30)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    run_main = title_p.add_run("VisionInspect AI")
    run_main.font.size = Pt(28)
    run_main.bold = True
    run_main.font.color.rgb = RGBColor(0, 120, 212)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(18)
    run_sub = subtitle_p.add_run("Milestone 2 Comprehensive Technical & Theoretical Documentation\nImage Processing, Quality Control, Defect Detection Engine, Severity Scoring & Batch Analytics Workflows")
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = RGBColor(100, 110, 120)
    
    # Meta table
    meta_tbl = doc.add_table(rows=2, cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_tbl.autofit = False
    
    meta_data = [
        [("Project / Platform:", " VisionInspect AI (Manufacturing Quality Inspection)"), ("Author / Developer:", " Ragul R V (AI/ML Engineering Intern)")],
        [("Milestone Covered:", " Milestone 2 (Stages 1, 2, 3 & 4)"), ("Document Version:", " 2.0 (Comprehensive Technical Reference)")]
    ]
    
    for r_idx, row in enumerate(meta_data):
        for c_idx, (lbl, val) in enumerate(row):
            cell = meta_tbl.cell(r_idx, c_idx)
            set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run1 = p.add_run(lbl)
            run1.bold = True
            run1.font.size = Pt(9.5)
            run1.font.color.rgb = RGBColor(70, 80, 95)
            run2 = p.add_run(val)
            run2.font.size = Pt(9.5)
            run2.font.color.rgb = RGBColor(30, 40, 50)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Divider Line
    div_p = doc.add_paragraph()
    div_p.paragraph_format.space_after = Pt(12)
    run_div = div_p.add_run("━" * 55)
    run_div.font.color.rgb = RGBColor(200, 210, 225)
    div_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # SECTION 1: EXECUTIVE SUMMARY
    h1 = doc.add_heading("1. Executive Summary & Objective", level=1)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "VisionInspect AI is an Industry 4.0 compliant automated computer vision platform designed for real-time quality control and manufacturing defect inspection. "
        "Milestone 2 focuses on building the core Image Processing Pipelines, Defect Detection Engines, Mathematical Severity Assessment Framework, Batch Analytics, and Interactive Dashboard UI."
    )
    
    create_callout(
        doc,
        "Milestone 2 transitions the system from basic baseline model inference to an end-to-end industrial inspection engine equipped with automated image validation (blur, lighting, contrast), morphological severity scoring, thread-safe batch processing, and live operational dashboards.",
        title="MILESTONE 2 CORE GOAL",
        fill_hex="EBF5FF",
        border_hex="0078D4"
    )

    # SECTION 2: ARCHITECTURE BLUEPRINT & STAGED ROADMAP
    h1 = doc.add_heading("2. System Pipeline Architecture & Staged Roadmap", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.add_run("To ensure modular reliability and zero-regression deployment, Milestone 2 was structured into 4 distinct execution stages:")

    roadmap_data = [
        ("Stage 1: Preprocessing & Quality Analysis", "Implements automated image validation metrics (sharpness/blur via Laplacian variance, exposure/brightness, contrast index) and enhancement filters (CLAHE in LAB color space, Gaussian denoising, unsharp masking). Adds the /quality-check API endpoint."),
        ("Stage 2: Defect Prediction Engine & Severity Scoring", "Integrates preprocessing with unsupervised Autoencoder models. Formulates the spec-compliant mathematical severity formula (Size 30% + Location 25% + Type 25% + Confidence 20%). Aligns model inference inputs with raw training data distributions."),
        ("Stage 3: Batch Inspection, History & Reporting API", "Implements a thread-safe in-memory FIFO log (inspection_log.py), structured Markdown/HTML/JSON quality certificates (report.py), and REST endpoints (/batch-predict, /history, /analytics, /report/{id})."),
        ("Stage 4: Enhanced Dashboard UI & Visual Analytics", "Upgrades dashboard.html into a multi-tab visual workstation featuring live severity gauges, single image QC meters, multi-file batch upload cards, Chart.js pass/fail ratios, and historical inspection tables.")
    ]

    for title_stage, desc_stage in roadmap_data:
        p_item = doc.add_paragraph(style='List Bullet')
        p_item.paragraph_format.space_after = Pt(4)
        run_t = p_item.add_run(f"{title_stage}: ")
        run_t.bold = True
        run_t.font.color.rgb = RGBColor(0, 100, 180)
        run_d = p_item.add_run(desc_stage)
        run_d.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Infographic / Flowchart diagram
    h2 = doc.add_heading("System Architecture & Data Processing Flowchart", level=2)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)

    p_flow = doc.add_paragraph()
    p_flow.paragraph_format.space_after = Pt(10)
    p_flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    flow_art = (
        "┌────────────────────────────────────────────────────────────────────────────────────────┐\n"
        "│                             PRODUCT IMAGE INPUT (Single / Batch)                       │\n"
        "└───────────────────────────────────────────┬────────────────────────────────────────────┘\n"
        "                                            │\n"
        "                                            ▼\n"
        "┌────────────────────────────────────────────────────────────────────────────────────────┐\n"
        "│ STAGE 1: IMAGE PROCESSING & QUALITY CONTROL (anomaly_detection/preprocessor.py)        │\n"
        "├────────────────────────────────────────────────────────────────────────────────────────┤\n"
        "│  1. Sharpness Metric : Laplacian Variance σ²(∇²f)  -->  Blur Warning (< 100.0)             │\n"
        "│  2. Exposure Check   : Mean Gray Intensity μ       -->  Dark (< 40) / Bright (> 230)       │\n"
        "│  3. Contrast Index   : Intensity Std Dev σ         -->  Low Contrast Warning (< 15.0)    │\n"
        "│  4. Enhancement      : CLAHE (LAB space L-channel) + Gaussian Denoise + Unsharp Mask   │\n"
        "└───────────────────────────────────────────┬────────────────────────────────────────────┘\n"
        "                                            │\n"
        "                      ┌─────────────────────┴─────────────────────┐\n"
        "                      │                                           │\n"
        "                      ▼                                           ▼\n"
        "┌──────────────────────────────────────────┐  ┌──────────────────────────────────────────┐\n"
        "│ STAGE 1 (YOLOv8): Product Isolation Crop │  │ QUALITY ANALYSIS REPORT (q_report)       │\n"
        "│ (Extract ROI, Bounding Box Coordinates)  │  │ Valid Status, Blur/Brightness Scores     │\n"
        "└─────────────────────┬────────────────────┘  └─────────────────────┬────────────────────┘\n"
        "                      │                                           │\n"
        "                      └─────────────────────┬─────────────────────┘\n"
        "                                            │\n"
        "                                            ▼\n"
        "┌────────────────────────────────────────────────────────────────────────────────────────┐\n"
        "│ STAGE 2: UNSUPERVISED AUTOENCODER RECONSTRUCTION (anomaly_detection/model.py)          │\n"
        "├────────────────────────────────────────────────────────────────────────────────────────┤\n"
        "│  - Input: Raw Cropped Image (Matches Training Distribution, Eliminates Shift)          │\n"
        "│  - Latent Bottleneck: 128x128x3 --> Conv Encoder --> 8x8x256 --> ConvTranspose Decoder │\n"
        "│  - Anomaly Map E(x): Pixel-wise Squared Error (x - x_hat)²                              │\n"
        "│  - Anomaly Score: Mean of Top 5% Highest Error Pixels                                 │\n"
        "│  - Classification Decision: Is Anomaly = (Score > Calibrated Category Threshold)       │\n"
        "└───────────────────────────────────────────┬────────────────────────────────────────────┘\n"
        "                                            │\n"
        "                                            ▼\n"
        "┌────────────────────────────────────────────────────────────────────────────────────────┐\n"
        "│ STAGE 2: SEVERITY SCORING ENGINE (anomaly_detection/severity.py)                       │\n"
        "├────────────────────────────────────────────────────────────────────────────────────────┤\n"
        "│  Severity Score = (Size x 30%) + (Location x 25%) + (Defect Type x 25%) + (Conf x 20%) │\n"
        "│  Levels: Critical (80-100) | High (60-79) | Medium (40-59) | Low (0-39)                  │\n"
        "└───────────────────────────────────────────┬────────────────────────────────────────────┘\n"
        "                                            │\n"
        "                                            ▼\n"
        "┌────────────────────────────────────────────────────────────────────────────────────────┐\n"
        "│ STAGE 3: BATCH LOGGING & REPORT GENERATION (inspection_log.py & report.py)             │\n"
        "├────────────────────────────────────────────────────────────────────────────────────────┤\n"
        "│  - Thread-Safe FIFO Log (Max 500 Entries)                                              │\n"
        "│  - Rest API Endpoints: /batch-predict, /history, /analytics, /report/{id}              │\n"
        "│  - Dynamic Certificate Formats: JSON, Raw Markdown, Styled Printable HTML              │\n"
        "└───────────────────────────────────────────┬────────────────────────────────────────────┘\n"
        "                                            │\n"
        "                                            ▼\n"
        "┌────────────────────────────────────────────────────────────────────────────────────────┐\n"
        "│ STAGE 4: WORKSTATION DASHBOARD UI (anomaly_detection/dashboard.html)                   │\n"
        "├────────────────────────────────────────────────────────────────────────────────────────┤\n"
        "│  - Multi-Tab UI: Single Inspect | Batch Inspect | Live Analytics | Inspection History   │\n"
        "│  - Visual Gauge Components, Heatmap Overlay, Chart.js Ratios & Certificate Modals     │\n"
        "└────────────────────────────────────────────────────────────────────────────────────────┘"
    )
    
    run_flow = p_flow.add_run(flow_art)
    run_flow.font.name = 'Consolas'
    run_flow.font.size = Pt(7.5)
    run_flow.font.color.rgb = RGBColor(0, 50, 100)

    # SECTION 3: STAGE 1 DEEP DIVE
    h1 = doc.add_heading("3. Stage 1: Preprocessing & Image Quality Validation", level=1)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "In industrial manufacturing lines, camera vibration, dirty lenses, conveyor speed, and uneven lighting often introduce out-of-focus blur or extreme exposure variance. "
        "Feeding unvalidated, degraded images into deep neural networks causes unpredictable reconstruction anomalies. "
        "Stage 1 establishes a deterministic image quality control filter and enhancement module in "
    )
    p.add_run("preprocessor.py").bold = True
    p.add_run(" to validate images prior to neural net evaluation.")

    h2 = doc.add_heading("3.1 Mathematical & Theoretical Foundation of Quality Metrics", level=2)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)

    # Quality metrics table
    q_table = doc.add_table(rows=4, cols=4)
    q_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    q_table.autofit = False

    headers = ["Metric Name", "Mathematical Expression", "Default Threshold", "Operational Meaning & Action"]
    hdr_row = q_table.rows[0]
    for idx, text in enumerate(headers):
        cell = hdr_row.cells[idx]
        set_cell_background(cell, "0078D4")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(255, 255, 255)

    q_data = [
        ("Sharpness (Blur)", "σ²(∇²f) = Var(Laplacian(Gray))", "Score < 100.0", "Detects high-frequency loss caused by camera focus errors or motion blur. Flags operator warning."),
        ("Brightness (Exposure)", "μ = (1 / HW) Σ x_ij", "40.0 <= μ <= 230.0", "Evaluates illumination intensity. Rejects under-exposed (dark) or over-exposed (saturated) images."),
        ("Contrast Index", "σ = sqrt( (1/HW) Σ (x_ij - μ)² )", "σ >= 15.0", "Measures pixel intensity dispersion. Rejects flat, washed-out images with low dynamic range.")
    ]

    for r_idx, row in enumerate(q_data):
        row_cells = q_table.rows[r_idx + 1].cells
        bg = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = row_cells[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            if c_idx == 0:
                run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    h2 = doc.add_heading("3.2 Image Enhancement Transformation Pipeline", level=2)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run("Images passing validation undergo three sequential visual enhancement filters:")

    enh_bullets = [
        ("CLAHE in LAB Color Space: ", "Standard histogram equalization on RGB channels causes severe color distortion and hue shifting. The preprocessor converts RGB to CIELAB color space, extracts the L* (Lightness) channel, applies Contrast Limited Adaptive Histogram Equalization with a clip limit of 2.0 and tile grid size of (8x8), and recombines the channels. This boosts local surface contrast while preserving true colors."),
        ("Gaussian Denoising: ", "Applies a mild 3x3 Gaussian kernel to suppress high-frequency electronic sensor grain without blurring major structural boundaries."),
        ("Unsharp Masking (Sharpening): ", "High-frequency edge details are restored using unsharp masking: Enhanced = 1.5 * Denoised - 0.5 * Blurred. This highlights micro-scratches and fine surface cracks for visual inspection.")
    ]

    for b_title, b_desc in enh_bullets:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(4)
        r1 = p_b.add_run(b_title)
        r1.bold = True
        r1.font.color.rgb = RGBColor(0, 120, 212)
        r2 = p_b.add_run(b_desc)
        r2.font.size = Pt(10)

    # Code snippet box for Preprocessor
    p_code_hdr = doc.add_paragraph()
    p_code_hdr.paragraph_format.space_before = Pt(8)
    p_code_hdr.paragraph_format.space_after = Pt(2)
    r_ch = p_code_hdr.add_run("Listing 1: Core Preprocessing & Quality Validation Implementation (preprocessor.py)")
    r_ch.bold = True
    r_ch.font.size = Pt(9.5)
    r_ch.font.color.rgb = RGBColor(80, 80, 80)

    p_code = doc.add_paragraph()
    p_code.paragraph_format.space_after = Pt(10)
    code_txt = (
        "def validate_and_preprocess_image(pil_img, blur_threshold=100.0, min_brightness=40.0, max_brightness=230.0, min_contrast=15.0):\
"
        "    img_bgr = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)\n"
        "    img_gray = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)\n"
        "    \n"
        "    # 1. Quality Metrics Computation\n"
        "    blur_score = float(cv2.Laplacian(img_gray, cv2.CV_64F).var())\n"
        "    brightness = float(np.mean(img_gray))\n"
        "    contrast = float(np.std(img_gray))\n"
        "    \n"
        "    warnings = []\n"
        "    is_valid = True\n"
        "    if blur_score < blur_threshold:\n"
        "        warnings.append(f'Low sharpness detected (Score: {blur_score:.1f} < {blur_threshold})')\n"
        "    if brightness < min_brightness or brightness > max_brightness:\n"
        "        is_valid = False\n"
        "    if contrast < min_contrast:\n"
        "        is_valid = False\n"
        "        \n"
        "    # 2. LAB CLAHE + Gaussian Blur + Unsharp Masking\n"
        "    lab = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2LAB)\n"
        "    l, a, b = cv2.split(lab)\n"
        "    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))\n"
        "    cl = clahe.apply(l)\n"
        "    enhanced_bgr = cv2.cvtColor(cv2.merge((cl, a, b)), cv2.COLOR_LAB2BGR)\n"
        "    denoised = cv2.GaussianBlur(enhanced_bgr, (3,3), 0)\n"
        "    sharpened = cv2.addWeighted(denoised, 1.5, cv2.GaussianBlur(denoised, (5,5), 0), -0.5, 0)\n"
        "    \n"
        "    return Image.fromarray(cv2.cvtColor(sharpened, cv2.COLOR_BGR2RGB)), {\n"
        "        'blur_score': round(blur_score, 2), 'brightness': round(brightness, 2),\n"
        "        'contrast': round(contrast, 2), 'is_valid': is_valid, 'warnings': warnings\n"
        "    }"
    )
    r_code = p_code.add_run(code_txt)
    r_code.font.name = 'Consolas'
    r_code.font.size = Pt(8.5)
    r_code.font.color.rgb = RGBColor(20, 30, 40)

    # SECTION 4: STAGE 2 DEEP DIVE
    h1 = doc.add_heading("4. Stage 2: Defect Detection Engine & Severity Scoring", level=1)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "Stage 2 integrates neural network reconstruction with a robust mathematical severity scoring framework. "
        "It addresses the core challenges of unsupervised anomaly localization, distribution shift prevention, and automated defect severity classification."
    )

    h2 = doc.add_heading("4.1 Convolutional Autoencoder Architecture & Anomaly Mapping", level=2)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "The system employs a Convolutional Autoencoder (AnomalyAutoencoder in model.py) trained exclusively on normal (good) product images. "
        "Because the model learns to compress and reconstruct normal textures, when presented with a defective product image containing cracks or scratches, "
        "it fails to reconstruct the unknown anomaly, creating a high pixel-wise reconstruction error map:"
    )

    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_eq.paragraph_format.space_after = Pt(6)
    r_eq = p_eq.add_run("E(x, y) = (1 / C) Σ_c ( x(x, y, c) - x̂(x, y, c) )²")
    r_eq.font.name = 'Consolas'
    r_eq.bold = True
    r_eq.font.size = Pt(11)
    r_eq.font.color.rgb = RGBColor(0, 100, 180)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "To prevent background noise from diluting small localized defects (e.g. hairline scratches), the overall scalar "
    )
    p.add_run("Anomaly Score").bold = True
    p.add_run(" is computed as the average of the top 5% highest pixel reconstruction errors:")

    p_eq2 = doc.add_paragraph()
    p_eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_eq2.paragraph_format.space_after = Pt(8)
    r_eq2 = p_eq2.add_run("Anomaly Score = (1 / K) Σ_{k=1}^{K} TopK( E(x, y), K = 0.05 · H · W )")
    r_eq2.font.name = 'Consolas'
    r_eq2.bold = True
    r_eq2.font.size = Pt(10.5)
    r_eq2.font.color.rgb = RGBColor(0, 100, 180)

    create_callout(
        doc,
        "Data Shift Resolution: Passing enhanced/sharpened images into the Autoencoder shifts pixel distributions and causes false positive anomaly spikes (e.g. 0.045 vs 0.017 threshold). In Stage 2, model inference is performed directly on the raw cropped image (matching PyTorch dataset training transforms), while the enhanced image is used strictly for operator visual inspection and quality metrics.",
        title="CRITICAL ENGINEERING INSIGHT",
        fill_hex="FFFBEB",
        border_hex="FF8C00"
    )

    h2 = doc.add_heading("4.2 Specification Severity Scoring Model", level=2)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run("Per the project specification, detected defects are assigned a composite Severity Score [0 - 100] calculated via:")

    p_sev_formula = doc.add_paragraph()
    p_sev_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sev_formula.paragraph_format.space_after = Pt(10)
    r_sf = p_sev_formula.add_run("Severity Score = (Size × 30%) + (Location × 25%) + (Defect Type × 25%) + (Confidence × 20%)")
    r_sf.font.name = 'Segoe UI'
    r_sf.bold = True
    r_sf.font.size = Pt(11)
    r_sf.font.color.rgb = RGBColor(180, 40, 0)

    # Severity Component breakdown table
    sev_tbl = doc.add_table(rows=5, cols=3)
    sev_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    sev_tbl.autofit = False

    s_headers = ["Parameter & Weight", "Mathematical / Morphological Calculation", "Industrial Significance"]
    s_hdr_row = sev_tbl.rows[0]
    for idx, text in enumerate(s_headers):
        cell = s_hdr_row.cells[idx]
        set_cell_background(cell, "D83B01")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(255, 255, 255)

    s_data = [
        ("Defect Size (30%)", "Ratio = Count(E > 0.5·Threshold) / (H·W)\nSize Score = min(100.0, (Ratio / 0.15) · 100)", "Quantifies physical surface area affected by the defect. A 15% surface defect receives 100% size score."),
        ("Defect Location (25%)", "Centroid of top 2% error pixels: (x_c, y_c)\nLocation Score = max(0, 100 - (Dist(Centroid, Center) / MaxDist)·100)", "Center-weighted. Defects occurring in critical central functional zones score higher than edge defects."),
        ("Defect Type (25%)", "Morphological Binarization + Connected Components:\nAspect Ratio > 3.0 --> Scratch (40)\nBlob > 5% Area --> Crack (95)\nDefault --> Contamination (70)", "Categorizes severity based on defect morphology. Structural cracks score higher than minor cosmetic scratches."),
        ("Detection Confidence (20%)", "If Score > Thresh: Conf = min(100, 80 + (Score/Thresh - 1)·40)\nIf Score <= Thresh: Conf = min(100, (1 - Score/Thresh)·100)", "Evaluates model certainty based on the scalar error margin relative to the calibrated decision threshold.")
    ]

    for r_idx, row in enumerate(s_data):
        row_cells = sev_tbl.rows[r_idx + 1].cells
        bg = "FFF8F6" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = row_cells[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            if c_idx == 0:
                run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Severity Level Mapping Table
    h3 = doc.add_heading("Severity Level & Action Mapping Rules", level=3)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)

    lvl_tbl = doc.add_table(rows=5, cols=3)
    lvl_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    lvl_tbl.autofit = False

    l_headers = ["Severity Range", "Level Classification", "Automated Action Trigger"]
    l_hdr_row = lvl_tbl.rows[0]
    for idx, text in enumerate(l_headers):
        cell = l_hdr_row.cells[idx]
        set_cell_background(cell, "1F2937")
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(255, 255, 255)

    l_data = [
        ("80.0 – 100.0", "Critical", "Reject Product & Trigger Quality Escalation Workflow"),
        ("60.0 – 79.9", "High", "Significant Quality Issue – Rework or Repair Recommended"),
        ("40.0 – 59.9", "Medium", "Moderate Concern – Requires QA Team Manual Review"),
        ("0.0 – 39.9", "Low", "Pass – Product is Generally Acceptable (Minor Cosmetic)")
    ]

    for r_idx, row in enumerate(l_data):
        row_cells = lvl_tbl.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cell = row_cells[c_idx]
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            if c_idx == 1:
                run.bold = True
                if val == "Critical": run.font.color.rgb = RGBColor(216, 59, 1)
                elif val == "High": run.font.color.rgb = RGBColor(255, 140, 0)
                elif val == "Medium": run.font.color.rgb = RGBColor(180, 140, 0)
                elif val == "Low": run.font.color.rgb = RGBColor(16, 124, 16)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    h2 = doc.add_heading("4.3 Calibrated Category Thresholds (3-Sigma Method)", level=2)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "Rather than using a fixed static threshold, calibrate_thresholds.py calculates thresholds for all 15 MVTec AD categories using a 3-Sigma statistical model: "
        "Threshold = Mean(Good) + max(3·Std(Good), 0.25·Mean(Good)). Below is the calibrated active threshold lookup implemented in api.py:"
    )

    t_tbl = doc.add_table(rows=4, cols=4)
    t_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_tbl.autofit = False

    t_headers = ["Category", "Calibrated Threshold", "Category", "Calibrated Threshold"]
    t_hdr_row = t_tbl.rows[0]
    for idx, text in enumerate(t_headers):
        cell = t_hdr_row.cells[idx]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)

    thresh_pairs = [
        ("bottle", "0.017216", "metal_nut", "0.019792"),
        ("cable", "0.028136", "pill", "0.005162"),
        ("capsule", "0.005667", "screw", "0.005689"),
        ("carpet", "0.014516", "tile", "0.016959"),
        ("grid", "0.011530", "toothbrush", "0.066125"),
        ("hazelnut", "0.004904", "transistor", "0.016376"),
        ("leather", "0.003908", "wood", "0.007200"),
        ("zipper", "0.010053", "-", "-")
    ]

    # Create rows dynamically
    for r_idx, (c1, t1, c2, t2) in enumerate(thresh_pairs):
        row_cells = t_tbl.add_row().cells
        bg = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate([c1, t1, c2, t2]):
            cell = row_cells[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=40, bottom=40, left=80, right=80)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(8.5)
            if c_idx in (1, 3) and val != "-":
                run.font.name = 'Consolas'
                run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # SECTION 5: STAGE 3 DEEP DIVE
    h1 = doc.add_heading("5. Stage 3: Batch Inspection, History & Reporting API", level=1)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "Stage 3 transitions the API into a multi-file batch processing service with thread-safe logging and automated report generation."
    )

    h2 = doc.add_heading("5.1 Thread-Safe Log & Inspection Certificates", level=2)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run("ThreadSafeInspectionLog (inspection_log.py): ")
    r1.bold = True
    p.add_run("Uses a Python threading.Lock to manage a FIFO memory queue capped at 500 entries. Stores full metadata including UUID, timestamp, quality metrics, anomaly scores, and severity breakdowns.")

    p2 = doc.add_paragraph(style='List Bullet')
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run("Report Certificate Generator (report.py): ")
    r2.bold = True
    p2.add_run("Formats inspection records into official Quality Inspection Certificates. Supports JSON payloads, raw Markdown text, and printable HTML documents with styled CSS status badges.")

    h2 = doc.add_heading("5.2 Stage 3 REST API Endpoints Specification", level=2)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)

    api_tbl = doc.add_table(rows=5, cols=4)
    api_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    api_tbl.autofit = False

    a_headers = ["Endpoint URL", "Method", "Request Payload", "Description & Response"]
    a_hdr_row = api_tbl.rows[0]
    for idx, text in enumerate(a_headers):
        cell = a_hdr_row.cells[idx]
        set_cell_background(cell, "0078D4")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(255, 255, 255)

    a_data = [
        ("/batch-predict", "POST", "files: List[UploadFile] (max 20)\ncategory: str, enable_yolo: bool", "Runs inspection on up to 20 images concurrently. Returns batch summary (pass rate, defect counts) and individual results."),
        ("/history", "GET", "limit: int (default 50, max 500)", "Retrieves recent inspection records from the thread-safe FIFO log in reverse chronological order."),
        ("/analytics", "GET", "None", "Returns real-time aggregated metrics: total inspections, pass/defect ratios, category statistics, and severity distributions."),
        ("/report/{inspection_id}", "GET", "format: json | markdown | html", "Fetches an inspection record by UUID and renders a formatted Quality Certificate in the requested format.")
    ]

    for r_idx, row in enumerate(a_data):
        row_cells = api_tbl.rows[r_idx + 1].cells
        bg = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = row_cells[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(8.5)
            if c_idx == 0:
                run.font.name = 'Consolas'
                run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # SECTION 6: STAGE 4 DEEP DIVE
    h1 = doc.add_heading("6. Stage 4: Workstation Dashboard UI", level=1)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "Stage 4 upgrades dashboard.html into an interactive, multi-tab workstation dashboard. "
        "Built using pure HTML5, CSS3 glassmorphism design tokens, Google Outfit typography, and Chart.js."
    )

    h2 = doc.add_heading("6.1 Multi-Tab Navigation Structure", level=2)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)

    tab_bullets = [
        ("Tab 1: Single Product Inspection: ", "Main single-image workflow. Renders the Pass/Fail decision badge, Stage 1 Image Quality meters (Sharpness, Exposure, Contrast), circular animated Severity Gauge, component breakdowns, and 4 visual cards (Original, Crop ROI, Model Reconstruction, Heatmap Overlay)."),
        ("Tab 2: Batch Inspection: ", "Multi-file drag-and-drop workstation. Allows uploading up to 20 images at once, displays real-time status pills (Queued, Processing, Success), summary cards, and individual result cards with direct links to view certificates."),
        ("Tab 3: Quality Analytics: ", "Real-time Chart.js dashboard displaying an approved vs rejected Pie Chart and a Severity Level Distribution Bar Chart."),
        ("Tab 4: Inspection Log: ", "Interactive history table fetching logs from /history, displaying timestamps, categories, scores, severity levels, and buttons to open HTML certificates in a new window.")
    ]

    for b_t, b_d in tab_bullets:
        p_t = doc.add_paragraph(style='List Bullet')
        p_t.paragraph_format.space_after = Pt(4)
        r1 = p_t.add_run(b_t)
        r1.bold = True
        r1.font.color.rgb = RGBColor(0, 120, 212)
        r2 = p_t.add_run(b_d)
        r2.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # SECTION 7: VERIFICATION & CODE APPENDIX
    h1 = doc.add_heading("7. Verification, Testing & Summary", level=1)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.add_run("All 4 stages of Milestone 2 were systematically verified using automated python test scripts:")

    test_table = doc.add_table(rows=5, cols=3)
    test_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    test_table.autofit = False

    th_headers = ["Verification Script", "Target Component", "Validation Result & Status"]
    th_hdr_row = test_table.rows[0]
    for idx, text in enumerate(th_headers):
        cell = th_hdr_row.cells[idx]
        set_cell_background(cell, "107C10")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(255, 255, 255)

    test_data = [
        ("verify_stage1.py", "preprocessor.py & /quality-check", "PASSED. Tested normal, blurry, dark, and overexposed image profiles."),
        ("verify_stage2.py", "severity.py & /predict pipeline", "PASSED. Verified normal vs anomalous score separation on bottle dataset."),
        ("verify_stage3.py", "inspection_log.py & /batch-predict", "PASSED. Tested batch processing, FIFO logs, analytics, and HTML/MD certificates."),
        ("test_pipeline.py", "Full Integrated Pipeline", "PASSED. Confirmed backward compatibility and full system health.")
    ]

    for r_idx, row in enumerate(test_data):
        row_cells = test_table.rows[r_idx + 1].cells
        bg = "F0FDF4" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = row_cells[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(8.5)
            if c_idx == 0:
                run.font.name = 'Consolas'
                run.bold = True
            elif c_idx == 2:
                run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    create_callout(
        doc,
        "Milestone 2 is 100% complete, fully tested, and committed to git branch 'RagulRV' (Commit: 0f81552). The platform is ready for Milestone 3 (Defect Classification & Manufacturing Analytics Integration).",
        title="MILESTONE 2 STATUS: COMPLETED",
        fill_hex="EBF5FF",
        border_hex="107C10"
    )

    # Save Document
    output_path = r"e:\Infosys Internship - 2 months\VisionInspectAI_Ragul_Model-Training\VisionInspectAI\VisionInspectAI_Milestone2_Documentation.docx"
    doc.save(output_path)
    print(f"Document successfully created and saved to: {output_path}")

if __name__ == "__main__":
    build_docx()
